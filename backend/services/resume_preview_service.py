from __future__ import annotations

import hashlib
import html
import os
import re
import shutil
from pathlib import Path
from typing import List, Sequence

from PIL import Image, ImageDraw, ImageFont

try:
    import fitz  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    fitz = None

try:
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    Document = None


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".heic", ".heif"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
DOC_EXTENSIONS = {".doc"}

PAGE_WIDTH = 1240
PAGE_HEIGHT = 1754
PAGE_MARGIN = 72
LINE_HEIGHT = 34
TITLE_HEIGHT = 54


def classify_resume_kind(file_name_or_path: str) -> str:
    ext = Path(file_name_or_path or "").suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in PDF_EXTENSIONS:
        return "pdf"
    if ext in DOCX_EXTENSIONS:
        return "docx"
    if ext in DOC_EXTENSIONS:
        return "doc"
    return "other"


def get_resume_asset_root(file_path: str) -> Path:
    path = Path(file_path)
    digest = hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:12]
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", path.name).strip("._") or "resume"
    return path.parent / "_resume_assets" / f"{safe_name}__{digest}"


def get_resume_preview_dir(file_path: str) -> Path:
    return get_resume_asset_root(file_path) / "preview_pages"


def list_resume_preview_images(file_path: str) -> List[Path]:
    kind = classify_resume_kind(file_path)
    source_path = Path(file_path)
    if kind == "image" and source_path.exists():
        return [source_path]

    preview_dir = get_resume_preview_dir(file_path)
    if not preview_dir.exists():
        return []

    return sorted(
        [path for path in preview_dir.iterdir() if path.is_file() and path.suffix.lower() == ".png"],
        key=lambda path: path.name,
    )


def cleanup_resume_assets(file_path: str) -> None:
    if not str(file_path or "").strip():
        return

    source_path = Path(file_path)
    if source_path.exists():
        try:
            source_path.unlink()
        except IsADirectoryError:
            shutil.rmtree(source_path, ignore_errors=True)

    asset_root = get_resume_asset_root(file_path)
    if asset_root.exists():
        shutil.rmtree(asset_root, ignore_errors=True)


def ensure_resume_previews(file_path: str, file_name: str = "") -> List[Path]:
    existing = list_resume_preview_images(file_path)
    if existing:
        return existing

    source_path = Path(file_path)
    if not source_path.exists():
        return []

    kind = classify_resume_kind(file_name or file_path)
    if kind == "image":
        return [source_path]

    preview_dir = get_resume_preview_dir(file_path)
    preview_dir.mkdir(parents=True, exist_ok=True)

    if kind == "pdf":
        rendered = _render_pdf_previews(source_path, preview_dir)
        if rendered:
            return rendered

    if kind == "docx":
        rendered = _render_docx_previews(source_path, preview_dir)
        if rendered:
            return rendered

    fallback = _render_placeholder_preview(
        preview_dir / "page-001.png",
        title=file_name or source_path.name,
        subtitle=_fallback_subtitle_for_kind(kind),
        body=_safe_preview_excerpt(source_path, kind),
    )
    return [fallback] if fallback else []


def get_resume_preview_summary(file_path: str, file_name: str = "") -> dict:
    preview_paths = ensure_resume_previews(file_path, file_name)
    kind = classify_resume_kind(file_name or file_path)
    return {
        "file_kind": kind,
        "preview_page_count": len(preview_paths),
        "preview_paths": [str(path) for path in preview_paths],
        "has_preview": bool(preview_paths),
    }


def _render_pdf_previews(source_path: Path, preview_dir: Path) -> List[Path]:
    if fitz is None:
        return []

    rendered: List[Path] = []
    with fitz.open(source_path) as document:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            target_path = preview_dir / f"page-{page_index + 1:03d}.png"
            pixmap.save(target_path)
            rendered.append(target_path)
    return rendered


def _render_docx_previews(source_path: Path, preview_dir: Path) -> List[Path]:
    if Document is None:
        return []

    document = Document(source_path)
    blocks = _collect_docx_blocks(document)
    if not blocks:
        blocks = ["未提取到可预览的正文内容。"]

    return _render_text_pages(
        blocks=blocks,
        preview_dir=preview_dir,
        title=source_path.name,
        subtitle="Word 预览",
    )


def _collect_docx_blocks(document: "Document") -> List[str]:
    blocks: List[str] = []

    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell_text for cell_text in (_normalize_text(cell.text) for cell in row.cells) if cell_text
            )
            if row_text:
                blocks.append(row_text)

    return blocks


def _render_text_pages(blocks: Sequence[str], preview_dir: Path, title: str, subtitle: str) -> List[Path]:
    font_title = _load_font(38)
    font_subtitle = _load_font(24)
    font_body = _load_font(24)

    usable_width = PAGE_WIDTH - PAGE_MARGIN * 2
    usable_height = PAGE_HEIGHT - PAGE_MARGIN * 2 - TITLE_HEIGHT - 48

    pages: List[Path] = []
    current_page_lines: List[str] = []
    current_height = 0

    for block in blocks:
        wrapped_lines = _wrap_text(block, font_body, usable_width)
        if not wrapped_lines:
            wrapped_lines = [""]

        block_height = len(wrapped_lines) * LINE_HEIGHT + 12
        if current_page_lines and current_height + block_height > usable_height:
            pages.append(
                _save_text_page(
                    preview_dir=preview_dir,
                    page_index=len(pages) + 1,
                    title=title,
                    subtitle=subtitle,
                    lines=current_page_lines,
                    font_title=font_title,
                    font_subtitle=font_subtitle,
                    font_body=font_body,
                )
            )
            current_page_lines = []
            current_height = 0

        current_page_lines.extend(wrapped_lines + [""])
        current_height += block_height

    if current_page_lines or not pages:
        pages.append(
            _save_text_page(
                preview_dir=preview_dir,
                page_index=len(pages) + 1,
                title=title,
                subtitle=subtitle,
                lines=current_page_lines or ["暂无内容"],
                font_title=font_title,
                font_subtitle=font_subtitle,
                font_body=font_body,
            )
        )

    return pages


def _save_text_page(
    preview_dir: Path,
    page_index: int,
    title: str,
    subtitle: str,
    lines: Sequence[str],
    font_title: ImageFont.ImageFont,
    font_subtitle: ImageFont.ImageFont,
    font_body: ImageFont.ImageFont,
) -> Path:
    canvas = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), "white")
    draw = ImageDraw.Draw(canvas)

    y = PAGE_MARGIN
    draw.text((PAGE_MARGIN, y), title, fill="#0f172a", font=font_title)
    y += TITLE_HEIGHT
    draw.text((PAGE_MARGIN, y), subtitle, fill="#475569", font=font_subtitle)
    y += 54

    for line in lines:
        draw.text((PAGE_MARGIN, y), line, fill="#111827", font=font_body)
        y += LINE_HEIGHT

    target_path = preview_dir / f"page-{page_index:03d}.png"
    canvas.save(target_path, format="PNG")
    return target_path


def _render_placeholder_preview(target_path: Path, title: str, subtitle: str, body: str) -> Path | None:
    target_path.parent.mkdir(parents=True, exist_ok=True)
    font_title = _load_font(34)
    font_subtitle = _load_font(24)
    font_body = _load_font(24)

    canvas = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), "white")
    draw = ImageDraw.Draw(canvas)
    draw.rounded_rectangle(
        (48, 48, PAGE_WIDTH - 48, PAGE_HEIGHT - 48),
        radius=28,
        outline="#dbe4f0",
        width=3,
        fill="#f8fafc",
    )

    y = 120
    draw.text((96, y), title, fill="#0f172a", font=font_title)
    y += 64
    draw.text((96, y), subtitle, fill="#475569", font=font_subtitle)
    y += 72

    for line in _wrap_text(body, font_body, PAGE_WIDTH - 192):
        draw.text((96, y), line, fill="#1f2937", font=font_body)
        y += LINE_HEIGHT
        if y > PAGE_HEIGHT - 120:
            break

    canvas.save(target_path, format="PNG")
    return target_path


def _safe_preview_excerpt(source_path: Path, kind: str) -> str:
    if kind == "docx" and Document is not None:
        try:
            blocks = _collect_docx_blocks(Document(source_path))
            return "\n".join(blocks[:12]) or "该简历已保存，请点击下载查看原件。"
        except Exception:
            pass
    return "该简历已保存到服务器。当前预览仅提供图片化查看，必要时可下载原件。"


def _fallback_subtitle_for_kind(kind: str) -> str:
    if kind == "doc":
        return "Word 文档预览"
    if kind == "other":
        return "文件预览"
    return "简历预览"


def _wrap_text(text: str, font: ImageFont.ImageFont, max_width: int) -> List[str]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    draw = ImageDraw.Draw(Image.new("RGB", (10, 10), "white"))
    lines: List[str] = []

    for paragraph in normalized.splitlines():
        paragraph = paragraph.strip()
        if not paragraph:
            lines.append("")
            continue

        current = ""
        for char in paragraph:
            candidate = current + char
            bbox = draw.textbbox((0, 0), candidate, font=font)
            width = bbox[2] - bbox[0]
            if current and width > max_width:
                lines.append(current)
                current = char
            else:
                current = candidate
        if current:
            lines.append(current)

    return lines


def _load_font(size: int) -> ImageFont.ImageFont:
    candidate_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
    ]
    for path in candidate_paths:
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def _normalize_text(value: object) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    text = html.unescape(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
